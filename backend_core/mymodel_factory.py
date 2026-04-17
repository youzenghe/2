from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

import markdown
from docx import Document
from flask import Flask, jsonify, redirect, render_template, request, send_file, url_for


@dataclass
class MyModelConfig:
    app_name: str
    docs_root: Path
    index_template: str = "myindex.html"
    show_template: str = "myshow.html"
    host: str = "127.0.0.1"
    port: int = 5029
    debug: bool = True


def create_mymodel_app(config: MyModelConfig) -> Flask:
    app = Flask(config.app_name)
    docs_root = config.docs_root.resolve()
    docs_root.mkdir(parents=True, exist_ok=True)
    allowed_extensions = {"md", "docx"}

    def secure_filename(filename: str | None) -> str:
        return Path(str(filename or "")).name

    def allowed_file(filename: str) -> bool:
        return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions

    def get_files() -> list[str]:
        files = [
            item.name
            for item in docs_root.iterdir()
            if item.is_file() and item.suffix.lower() in {".md", ".docx"}
        ]
        return sorted(files, key=str.lower)

    def resolve_doc_path(filename: str | None) -> tuple[Path, str]:
        safe_name = secure_filename(filename)
        if not safe_name or not allowed_file(safe_name):
            raise ValueError("Invalid file or file type not supported")
        target = (docs_root / safe_name).resolve()
        if target.parent != docs_root:
            raise ValueError("Invalid file path")
        return target, safe_name

    class HTMLTextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._chunks: list[str] = []

        def handle_data(self, data: str) -> None:
            self._chunks.append(data)

        def get_text(self) -> str:
            return "\n".join(self._chunks)

    def html_to_text(html: str) -> str:
        extractor = HTMLTextExtractor()
        extractor.feed(html)
        return extractor.get_text()

    @app.route("/")
    def index():
        return render_template(config.index_template, all_files=get_files())

    @app.route("/show_md", methods=["GET"])
    def show_md():
        filename = request.args.get("filename")
        try:
            file_path, safe_name = resolve_doc_path(filename)
        except ValueError as exc:
            return str(exc), 400

        if not file_path.exists():
            return "File not found", 404

        suffix = file_path.suffix.lower()
        if suffix == ".md":
            md_content = file_path.read_text(encoding="utf-8")
            content = markdown.markdown(md_content)
            return render_template(config.show_template, content=content, filename=safe_name, file_type="md")
        if suffix == ".docx":
            document = Document(str(file_path))
            content = "\n".join(paragraph.text for paragraph in document.paragraphs).replace("\n", "<br>")
            return render_template(config.show_template, content=content, filename=safe_name, file_type="docx")
        return "This file type is not supported for viewing", 400

    @app.route("/download", methods=["GET"])
    def download():
        filename = request.args.get("filename")
        try:
            file_path, _ = resolve_doc_path(filename)
        except ValueError:
            return "Filename not provided", 400
        if not file_path.exists():
            return "File not found", 404
        return send_file(str(file_path), as_attachment=True)

    @app.route("/save_md", methods=["POST"])
    def save_md():
        filename = request.form.get("filename")
        content = request.form.get("content")
        if not filename or content is None:
            return jsonify({"status": "error", "message": "Missing filename or content"}), 400
        try:
            file_path, _ = resolve_doc_path(filename)
            if file_path.suffix.lower() == ".docx":
                plain_text = html_to_text(content)
                doc = Document()
                for paragraph in plain_text.split("\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                doc.save(str(file_path))
            else:
                file_path.write_text(content, encoding="utf-8")
            return jsonify({"status": "success", "message": "File saved successfully"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    @app.route("/create", methods=["GET"])
    def create():
        filename = request.args.get("filename")
        try:
            file_path, safe_name = resolve_doc_path(filename)
        except ValueError as exc:
            return str(exc), 400
        if file_path.exists():
            return "File already exists", 400
        try:
            if file_path.suffix.lower() == ".docx":
                Document().save(str(file_path))
            else:
                file_path.write_text("", encoding="utf-8")
            return redirect(url_for("show_md", filename=safe_name))
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    @app.route("/rename_file", methods=["POST"])
    def rename_file():
        old_filename = request.form.get("old_filename")
        new_filename = request.form.get("new_filename")
        if not old_filename or not new_filename:
            return jsonify({"status": "error", "message": "Missing filenames"}), 400
        try:
            old_path, _ = resolve_doc_path(old_filename)
            new_path, _ = resolve_doc_path(new_filename)
        except ValueError as exc:
            return jsonify({"status": "error", "message": str(exc)}), 400
        if not old_path.exists():
            return jsonify({"status": "error", "message": "Old file not found"}), 404
        if new_path.exists():
            return jsonify({"status": "error", "message": "New file already exists"}), 400
        try:
            old_path.rename(new_path)
            return jsonify({"status": "success", "message": "File renamed successfully"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    @app.route("/delete_file", methods=["POST"])
    def delete_file():
        filename = request.form.get("filename")
        if not filename:
            return jsonify({"status": "error", "message": "Missing filename"}), 400
        try:
            file_path, _ = resolve_doc_path(filename)
        except ValueError as exc:
            return jsonify({"status": "error", "message": str(exc)}), 400
        if not file_path.exists():
            return jsonify({"status": "error", "message": "File not found"}), 404
        try:
            file_path.unlink()
            return jsonify({"status": "success", "message": "File deleted successfully"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    return app

