from __future__ import annotations

import json
import logging
import mimetypes
import os
from io import BytesIO
from pathlib import Path
from typing import Dict, List

import docx
import fitz
from fastapi import FastAPI
from fastapi import File
from fastapi import Form
from fastapi import Path as ApiPath
from fastapi import UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from starlette.responses import FileResponse
from starlette.responses import StreamingResponse

from .legal_agent_models import ComplianceResponse
from .legal_agent_models import ComplianceResponseData
from .legal_agent_models import ComplianceResult
from ..compliance_engine import ComplianceFinding, analyze_with_chunking, normalize_text
from ..deepseek_gateway import DeepSeekComplianceGateway, DeepSeekConfig

REPO_ROOT = Path(__file__).resolve().parents[2]

LEGAL_BACKEND_DIR = REPO_ROOT / "legal-agent" / "legal-agent-backend"

DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}
PDF_MIME_TYPES = {
    "application/pdf",
    "application/octet-stream",
}

RISK_KEYWORDS: Dict[str, set[str]] = {
    "payment": {"付款", "结算", "违约金", "滞纳金"},
    "liability": {"赔偿", "责任", "免责", "损失"},
    "delivery": {"交付", "验收", "质量", "履行"},
    "dispute": {"争议解决", "管辖", "仲裁", "法院"},
    "ip": {"知识产权", "保密", "商标", "专利"},
}

FOCUS_WEIGHTS: Dict[str, Dict[str, float]] = {
    "standard": {"payment": 1.0, "liability": 1.0, "delivery": 1.0, "dispute": 1.0, "ip": 1.0},
    "party_a": {"payment": 1.4, "liability": 1.2, "delivery": 1.2, "dispute": 1.2, "ip": 1.0},
    "party_b": {"payment": 1.0, "liability": 1.4, "delivery": 1.2, "dispute": 1.2, "ip": 1.0},
    "strict": {"payment": 1.5, "liability": 1.5, "delivery": 1.3, "dispute": 1.3, "ip": 1.3},
}

mimetypes.add_type("text/css", ".css")
mimetypes.add_type("application/javascript", ".js")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=LEGAL_BACKEND_DIR / "static"), name="static")
app.mount("/assets", StaticFiles(directory=LEGAL_BACKEND_DIR / "static" / "dist" / "assets"), name="assets")

_gateway = DeepSeekComplianceGateway(
    DeepSeekConfig(
        api_key=os.getenv("DEEPSEEK_API_KEY", "sk-a3f6222df2504e4fb102ac50c6b98980"),
        base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1"),
        model=os.getenv("DEEPSEEK_MODEL", "deepseek-chat"),
        timeout_seconds=int(os.getenv("DEEPSEEK_TIMEOUT_SECONDS", "80")),
        retry_times=int(os.getenv("DEEPSEEK_RETRY_TIMES", "2")),
    )
)


def _build_error_response(message: str, code: int = 400) -> ComplianceResponse:
    return ComplianceResponse(
        message=message,
        code=code,
        data=ComplianceResponseData(compliance=False, result=[]),
    )


def _file_suffix(file: UploadFile) -> str:
    return Path(file.filename or "").suffix.lower()


def _is_pdf_upload(file: UploadFile) -> bool:
    content_type = (file.content_type or "").lower()
    suffix = _file_suffix(file)
    return content_type in PDF_MIME_TYPES and suffix in {".pdf", ""} or suffix == ".pdf"


def _is_docx_upload(file: UploadFile) -> bool:
    content_type = (file.content_type or "").lower()
    suffix = _file_suffix(file)
    return content_type in DOCX_MIME_TYPES and suffix in {".docx", ""} or suffix == ".docx"


def _extract_contract_text(file: UploadFile, contents: bytes) -> str:
    if _is_pdf_upload(file):
        pdf_doc = fitz.open(stream=contents, filetype="pdf")
        try:
            pages = [page.get_text() for page in pdf_doc]
        finally:
            pdf_doc.close()
        return "\n".join(pages)

    if _is_docx_upload(file):
        document = docx.Document(BytesIO(contents))
        return "\n".join(para.text for para in document.paragraphs)

    raise ValueError("Unsupported file type")


def _result_key(item: ComplianceResult) -> str:
    return f"{item.original.strip()[:180]}::{item.suggestion.strip()[:180]}"


def _dedupe_model_results(items: List[ComplianceResult], max_items: int = 40) -> List[ComplianceResult]:
    deduped: List[ComplianceResult] = []
    seen = set()
    for item in items:
        key = _result_key(item)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
        if len(deduped) >= max_items:
            break
    return deduped


def _findings_to_models(findings: List[ComplianceFinding]) -> List[ComplianceResult]:
    models: List[ComplianceResult] = []
    for item in findings:
        try:
            models.append(ComplianceResult(original=item.original, risk=item.risk, suggestion=item.suggestion))
        except Exception as exc:  # noqa: BLE001
            logging.warning("[compliance] invalid finding skipped: %s", exc)
    return _dedupe_model_results(models)


def _deepseek_provider_call(contract_chunk: str, focus: str) -> str:
    return _gateway.review("No extra legal references provided.", contract_chunk, focus)


def _is_supported_provider(provider: str) -> bool:
    normalized_provider = (provider or "").strip().lower()
    return normalized_provider in {
        "deepseek",
        "openai",
        "tongyi",
        "qwen",
        "qianwen",
        "kimi",
        "oneapi",
        "ollama",
        "local",
        "local-lm",
        "local_lm",
    }


@app.get("/")
async def read_index():
    return FileResponse(LEGAL_BACKEND_DIR / "static" / "dist" / "index.html")


@app.get("/ping")
async def root():
    return "pong"


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    return {"message": "File uploaded", "filename": file.filename, "status": "done"}


def replace_text_in_paragraph(paragraph, original: str, suggestion: str) -> None:
    if original not in paragraph.text:
        return

    for run in paragraph.runs:
        if original in run.text:
            run.text = run.text.replace(original, suggestion)
            return

    full_text = paragraph.text
    new_text = full_text.replace(original, suggestion)
    base_run = next((run for run in paragraph.runs if run.text.strip()), None)

    for run in paragraph.runs:
        run.text = ""

    if base_run is not None:
        base_run.text = new_text
    elif paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


@app.post("/compliance/merge")
async def compliance_merge(file: UploadFile = File(...), replacements: str = Form(...)):
    if not _is_docx_upload(file):
        return _build_error_response("Only DOCX is supported for merge export.", code=400)

    contents = await file.read()
    document = docx.Document(BytesIO(contents))

    try:
        replacement_list = json.loads(replacements)
    except Exception:  # noqa: BLE001
        return _build_error_response("Invalid replacements payload.", code=400)

    if not isinstance(replacement_list, list):
        return _build_error_response("Replacements payload must be a list.", code=400)

    for item in replacement_list:
        if not isinstance(item, dict):
            continue
        original = str(item.get("original") or "")
        suggestion = str(item.get("suggestion") or "")
        if not original or not suggestion:
            continue

        for paragraph in document.paragraphs:
            replace_text_in_paragraph(paragraph, original, suggestion)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, original, suggestion)

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=merged_contract.docx"},
    )


@app.post("/compliance/{provider}")
async def compliance(provider: str = ApiPath(...), file: UploadFile = File(...), focus: str = Form("standard")):
    if not _is_supported_provider(provider):
        return _build_error_response("Invalid provider. Please provide a valid provider.", code=400)

    if not _is_docx_upload(file) and not _is_pdf_upload(file):
        return _build_error_response("Unsupported file type. Please upload DOCX or PDF.", code=400)

    contents = await file.read()
    try:
        contract_text = _extract_contract_text(file, contents)
        contract_text = normalize_text(contract_text)
    except Exception as exc:  # noqa: BLE001
        logging.exception("[compliance] extract text failed: %s", exc)
        return _build_error_response("Failed to parse file content.", code=400)

    if not contract_text:
        return _build_error_response("Contract content is empty.", code=400)

    logging.info(
        "[compliance] request provider=%s filename=%s focus=%s chars=%s",
        provider,
        file.filename,
        focus,
        len(contract_text),
    )

    findings = analyze_with_chunking(
        contract_text=contract_text,
        focus=focus,
        provider_call=_deepseek_provider_call,
        risk_keywords=RISK_KEYWORDS,
        focus_weights=FOCUS_WEIGHTS,
    )
    result_items = _findings_to_models(findings)
    logging.info("[compliance] response findings=%s filename=%s", len(result_items), file.filename)

    if not result_items:
        return _build_error_response("Compliance analysis unavailable, please retry later.", code=502)

    response_data = ComplianceResponseData(compliance=len(result_items) > 0, result=result_items)
    return ComplianceResponse(message="ok", code=200, data=response_data)
