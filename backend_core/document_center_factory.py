from __future__ import annotations

import difflib
import hashlib
import json
import re
import shutil
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from urllib.parse import quote

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from flask import Flask, jsonify, make_response, render_template, request, send_file
from markupsafe import Markup


@dataclass
class DocumentCenterConfig:
    app_name: str
    base_dir: Path
    docs_root: Path
    mymodel_root: Path
    versions_root: Path
    temp_root: Path
    favorites_file: Path
    recent_file: Path
    category_rules: Sequence[Tuple[str, Sequence[str]]]
    risk_keywords: Sequence[Tuple[str, str]]
    clause_templates: Dict[str, Dict[str, List[Dict[str, str]]]]
    default_clause_template: Dict[str, List[Dict[str, str]]]
    reminder_patterns: Sequence[Tuple[str, str]]
    index_template: str = "index.html"
    show_template: str = "show_md.html"
    host: str = "127.0.0.1"
    port: int = 5002
    debug: bool = True


def create_document_center_app(config: DocumentCenterConfig) -> Flask:
    app = Flask(config.app_name)
    app.secret_key = f"{config.app_name}-secret"

    docs_root = config.docs_root.resolve()
    mymodel_root = config.mymodel_root.resolve()
    versions_root = config.versions_root.resolve()
    temp_root = config.temp_root.resolve()
    favorites_file = config.favorites_file.resolve()
    recent_file = config.recent_file.resolve()

    for folder in (docs_root, mymodel_root, versions_root, temp_root):
        folder.mkdir(parents=True, exist_ok=True)

    allowed_roots = [docs_root, mymodel_root, versions_root, temp_root]
    reminders_db: List[Dict[str, Any]] = []
    reminder_id_counter = 1
    all_files: List[str] = []

    def is_relative_to(path: Path, root: Path) -> bool:
        try:
            path.relative_to(root)
            return True
        except ValueError:
            return False

    def is_allowed_path(path: Path) -> bool:
        resolved = path.resolve(strict=False)
        return any(is_relative_to(resolved, root) for root in allowed_roots)

    def load_json_list(path: Path) -> List[str]:
        if not path.exists():
            return []
        try:
            with path.open("r", encoding="utf-8") as handle:
                payload = json.load(handle)
            if isinstance(payload, list):
                return payload
        except Exception:
            pass
        return []

    def save_json(path: Path, data: Any) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", encoding="utf-8") as handle:
            json.dump(data, handle, ensure_ascii=False, indent=2)

    def safe_filename(filename: str, default_stem: str = "document", suffix: str = ".md") -> str:
        stem = Path(filename or "").name.strip()
        if not stem:
            stem = default_stem + suffix
        if not stem.lower().endswith(suffix.lower()):
            stem += suffix
        return stem

    def read_text(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_bytes().decode("utf-8", errors="ignore")

    def write_text(path: Path, content: str) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")

    def compact_path_text(value: str) -> str:
        return re.sub(r"[\\/:\s]+", "", str(value or "")).lower()

    def ascii_download_name(filename: str, default_stem: str = "download") -> str:
        path = Path(filename or "")
        stem = unicodedata.normalize("NFKD", path.stem).encode("ascii", "ignore").decode("ascii").strip(" ._-")
        if not stem:
            stem = default_stem
        return f"{stem}{path.suffix}"

    def build_content_disposition(filename: str, default_stem: str = "download") -> str:
        fallback_name = ascii_download_name(filename, default_stem=default_stem)
        encoded_name = quote(str(filename or ""), safe="")
        return f"attachment; filename=\"{fallback_name}\"; filename*=UTF-8''{encoded_name}"

    def resolve_existing_path(raw_path: str) -> Optional[Path]:
        if not raw_path:
            return None
        input_path = Path(raw_path)
        candidates: List[Path] = []
        if input_path.is_absolute():
            candidates.append(input_path)
        else:
            candidates.extend(
                [
                    (config.base_dir / input_path).resolve(strict=False),
                    (docs_root / input_path).resolve(strict=False),
                    (mymodel_root / input_path).resolve(strict=False),
                ]
            )
        for candidate in candidates:
            resolved = candidate.resolve(strict=False)
            if resolved.exists() and is_allowed_path(resolved):
                return resolved
        basename = input_path.name
        for root in (docs_root, mymodel_root):
            fallback = root / basename
            if fallback.exists() and is_allowed_path(fallback):
                return fallback.resolve(strict=False)
        compact_raw = compact_path_text(raw_path)
        extension_match = re.search(r"\.(md|docx|txt|pdf)$", str(raw_path), re.IGNORECASE)
        if compact_raw and extension_match:
            pattern = f"*.{extension_match.group(1).lower()}"
            basename_compact = compact_path_text(input_path.name)
            for root in allowed_roots:
                for candidate in root.rglob(pattern):
                    resolved = candidate.resolve(strict=False)
                    if not resolved.exists() or not is_allowed_path(resolved):
                        continue
                    candidate_compact = compact_path_text(str(resolved))
                    if compact_raw == candidate_compact:
                        return resolved
                    if basename_compact and basename_compact == compact_path_text(resolved.name):
                        return resolved
                    if compact_raw.endswith(candidate_compact):
                        return resolved
        return None

    def normalize_file_ref(path: Path) -> str:
        return str(path.resolve(strict=False))

    def next_available_path(target: Path) -> Path:
        if not target.exists():
            return target
        index = 1
        while True:
            candidate = target.with_name(f"{target.stem}_{index}{target.suffix}")
            if not candidate.exists():
                return candidate
            index += 1

    def markdown_to_document(md_content: str) -> Document:
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, "html.parser")
        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "blockquote"]):
            tag = element.name or ""
            text = element.get_text("\n").strip()
            if not text:
                continue
            if tag.startswith("h"):
                level = int(tag[1]) if len(tag) > 1 and tag[1].isdigit() else 1
                heading = document.add_heading(level=max(1, min(level, 6)))
                heading.add_run(text).bold = True
            elif tag == "p":
                document.add_paragraph(text)
            elif tag in {"ul", "ol"}:
                style_name = "List Bullet" if tag == "ul" else "List Number"
                for item in element.find_all("li", recursive=False):
                    document.add_paragraph(item.get_text(" ").strip(), style=style_name)
            elif tag == "blockquote":
                quote = document.add_paragraph(text)
                quote.style = "Intense Quote"
        return document

    def upload_to_text(upload) -> str:
        raw = upload.read()
        filename = (upload.filename or "").lower()
        if filename.endswith(".docx"):
            document = Document(BytesIO(raw))
            return "\n".join(p.text for p in document.paragraphs if p.text)
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore")

    @app.template_filter("basename")
    def basename_filter(path: str) -> str:
        return Path(path or "").name

    @app.template_filter("urlencode")
    def urlencode_filter(value: str) -> str:
        return quote(str(value or ""), safe="")

    def scan_files() -> List[str]:
        nonlocal all_files
        files: List[str] = []
        for root in (docs_root, mymodel_root):
            if root.exists():
                for file in root.rglob("*.md"):
                    files.append(normalize_file_ref(file))
        all_files = sorted(files)
        return all_files

    favorites = load_json_list(favorites_file)
    recent_files = load_json_list(recent_file)
    scan_files()

    def categorize_file(filename: str) -> str:
        name = Path(filename).name.lower()
        for category, tokens in config.category_rules:
            if any(token.lower() in name for token in tokens):
                return category
        if "合同" in Path(filename).name:
            return "其他合同"
        return "其他"

    def md_to_html(filename: str) -> Markup:
        file_path = resolve_existing_path(filename)
        if not file_path:
            return Markup("<p>File not found.</p>")
        return Markup(markdown.markdown(read_text(file_path)))

    def version_dir_for(file_path: Path) -> Path:
        digest = hashlib.sha1(str(file_path).encode("utf-8")).hexdigest()[:10]
        return (versions_root / f"{file_path.stem}_{digest}").resolve()

    def save_version(file_path: str, content: str) -> str:
        source = resolve_existing_path(file_path) or Path(file_path).resolve(strict=False)
        if not is_allowed_path(source):
            raise ValueError("Invalid file path")
        version_dir = version_dir_for(source)
        version_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        version_file = version_dir / f"{source.stem}_{timestamp}.md"
        write_text(version_file, content)
        return normalize_file_ref(version_file)

    def get_file_versions(file_path: str) -> List[Dict[str, Any]]:
        source = resolve_existing_path(file_path)
        if not source:
            return []
        version_dir = version_dir_for(source)
        if not version_dir.exists():
            return []
        versions = []
        for item in sorted(version_dir.glob("*.md"), reverse=True):
            stat = item.stat()
            versions.append(
                {
                    "filename": item.name,
                    "path": normalize_file_ref(item),
                    "time": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                    "size": stat.st_size,
                }
            )
        return versions

    def add_to_recent(filename: str) -> None:
        nonlocal recent_files
        if not filename:
            return
        if filename in recent_files:
            recent_files.remove(filename)
        recent_files.insert(0, filename)
        recent_files = recent_files[:10]
        save_json(recent_file, recent_files)

    def normalize_date(value: str) -> Optional[str]:
        raw = (value or "").strip()
        if not raw:
            return None
        attempts = [("%Y-%m-%d", raw), ("%Y/%m/%d", raw)]
        if "年" in raw and "月" in raw:
            normalized = raw.replace("年", "-").replace("月", "-").replace("日", "")
            attempts.append(("%Y-%m-%d", normalized))
        for fmt, candidate in attempts:
            try:
                return datetime.strptime(candidate, fmt).strftime("%Y-%m-%d")
            except ValueError:
                continue
        return None

    def extract_dates(text: str) -> List[Dict[str, str]]:
        extracted: List[Dict[str, str]] = []
        for pattern, event in config.reminder_patterns:
            for match in re.findall(pattern, text or ""):
                date_value = normalize_date(match)
                if date_value:
                    extracted.append({"date": date_value, "event": event})
        return extracted

    def analyze_legal_risks(
        content1: str,
        content2: str,
        lines1: Sequence[str],
        lines2: Sequence[str],
        matcher: difflib.SequenceMatcher,
    ) -> List[Dict[str, str]]:
        del content1, content2
        risks: List[Dict[str, str]] = []
        seen = set()
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag not in {"replace", "delete", "insert"}:
                continue
            changed_lines: List[str] = []
            if tag in {"replace", "delete"}:
                changed_lines.extend(lines1[i1:i2])
            if tag in {"replace", "insert"}:
                changed_lines.extend(lines2[j1:j2])
            changed_text = " ".join(changed_lines).lower()
            for keyword, message in config.risk_keywords:
                if keyword.lower() in changed_text and keyword not in seen:
                    seen.add(keyword)
                    risks.append({"clause": keyword, "message": message, "status": "modified"})
        return risks

    @app.route("/")
    def index():
        category_counter: Dict[str, int] = {}
        for file in all_files:
            category = categorize_file(file)
            category_counter[category] = category_counter.get(category, 0) + 1
        categories = [{"name": key, "count": value} for key, value in category_counter.items()]
        categories.sort(key=lambda item: item["count"], reverse=True)
        return render_template(
            config.index_template,
            all_files=all_files,
            categories=categories,
            total_files=len(all_files),
            favorites=favorites,
            recent_files=recent_files[:10],
        )

    @app.route("/show_md", methods=["GET"])
    def show_md():
        filename = request.args.get("filename", "")
        file_path = resolve_existing_path(filename)
        if not file_path:
            return "File not found", 404
        file_ref = normalize_file_ref(file_path)
        add_to_recent(file_ref)
        content = md_to_html(file_ref)
        md_content = read_text(file_path)
        versions = get_file_versions(file_ref)
        return render_template(
            config.show_template,
            content=content,
            md_content=md_content,
            filename=file_ref,
            display_filename=file_path.name,
            display_stem=file_path.stem,
            versions=versions,
        )

    @app.route("/download_word", methods=["GET"])
    def download_word():
        filename = request.args.get("filename", "")
        file_path = resolve_existing_path(filename)
        if not file_path:
            return "File not found", 404
        response = send_file(file_path, as_attachment=False)
        response.headers["Content-Disposition"] = build_content_disposition(file_path.name, default_stem="document")
        return response

    @app.route("/save_md", methods=["POST"])
    def save_md():
        custom_filename = request.form.get("custom_filename", "").strip()
        content = request.form.get("content", "")
        create_version = request.form.get("create_version", "false").lower() == "true"
        if not custom_filename:
            return jsonify({"status": "error", "message": "Invalid filename"}), 400
        safe_name = safe_filename(custom_filename, default_stem="custom", suffix=".md")
        target = (mymodel_root / safe_name).resolve(strict=False)
        if not is_relative_to(target, mymodel_root):
            return jsonify({"status": "error", "message": "Invalid path"}), 400
        try:
            if create_version and target.exists():
                save_version(str(target), read_text(target))
            write_text(target, content)
            scan_files()
            return jsonify({"status": "success", "message": "File saved successfully"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    @app.route("/delete_temp_file", methods=["POST"])
    def delete_temp_file():
        filename = request.form.get("filename", "")
        file_path = resolve_existing_path(filename)
        try:
            if file_path and file_path.exists():
                file_path.unlink()
            scan_files()
            return "File deleted successfully"
        except Exception as exc:  # noqa: BLE001
            return str(exc), 500

    @app.route("/convert_to_word", methods=["GET"])
    def convert_to_word():
        md_filename = request.args.get("filename", "")
        file_path = resolve_existing_path(md_filename)
        if not file_path:
            return "File not found", 404
        document = markdown_to_document(read_text(file_path))
        docx_io = BytesIO()
        document.save(docx_io)
        docx_io.seek(0)
        response = make_response(docx_io.getvalue())
        response.headers.set("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response.headers["Content-Disposition"] = build_content_disposition(f"{file_path.stem}.docx", default_stem="document")
        return response

    @app.route("/save_as_word", methods=["POST"])
    def save_as_word():
        custom_filename = request.form.get("custom_filename", "").strip()
        content = request.form.get("content", "")
        safe_name = Path(custom_filename).name or "custom"
        if safe_name.lower().endswith(".docx"):
            safe_name = Path(safe_name).stem
        target = (mymodel_root / f"{safe_name}.docx").resolve(strict=False)
        if not is_relative_to(target, mymodel_root):
            return jsonify({"status": "error", "message": "Invalid path"}), 400
        try:
            markdown_to_document(content).save(target)
            return jsonify({"status": "success", "message": "File saved successfully"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)}), 500

    @app.route("/get_versions", methods=["POST"])
    def get_versions():
        payload = request.get_json(silent=True) or {}
        filename = payload.get("filename", "")
        return jsonify({"status": "success", "versions": get_file_versions(filename)})

    @app.route("/compare_versions", methods=["POST"])
    def compare_versions():
        payload = request.get_json(silent=True) or {}
        version1 = resolve_existing_path(payload.get("version1", ""))
        version2 = resolve_existing_path(payload.get("version2", ""))
        if not version1 or not version2:
            return jsonify({"status": "error", "message": "Version file not found"})
        try:
            content1 = read_text(version1).splitlines(keepends=True)
            content2 = read_text(version2).splitlines(keepends=True)
            diff_html = difflib.HtmlDiff().make_file(content1, content2, context=True)
            return jsonify({"status": "success", "diff_html": diff_html})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)})

    @app.route("/restore_version", methods=["POST"])
    def restore_version():
        payload = request.get_json(silent=True) or {}
        version_path = resolve_existing_path(payload.get("version_path", ""))
        original_file = resolve_existing_path(payload.get("original_file", ""))
        if not version_path or not original_file:
            return jsonify({"status": "error", "message": "File not found"})
        try:
            if original_file.exists():
                save_version(str(original_file), read_text(original_file))
            write_text(original_file, read_text(version_path))
            scan_files()
            return jsonify({"status": "success", "message": "Version restored"})
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)})

    @app.route("/toggle_favorite", methods=["POST"])
    def toggle_favorite():
        payload = request.get_json(silent=True) or {}
        filename = payload.get("filename", "")
        if not filename:
            return jsonify({"status": "error", "message": "Invalid filename"}), 400
        if filename in favorites:
            favorites.remove(filename)
            message = "Favorite removed"
        else:
            favorites.append(filename)
            message = "Favorite added"
        save_json(favorites_file, favorites)
        return jsonify({"status": "success", "is_favorite": filename in favorites, "message": message})

    @app.route("/get_favorites", methods=["GET"])
    def get_favorites():
        return jsonify({"status": "success", "favorites": favorites})

    @app.route("/get_recent", methods=["GET"])
    def get_recent():
        return jsonify({"status": "success", "recent": recent_files[:10]})

    @app.route("/batch_download", methods=["POST"])
    def batch_download():
        payload = request.get_json(silent=True) or {}
        files = payload.get("files", [])
        if not files:
            return jsonify({"status": "error", "message": "Please select files"})
        zip_name = f"batch_download_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        zip_path = (temp_root / zip_name).resolve()
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for item in files:
                file_path = resolve_existing_path(item)
                if file_path and file_path.exists():
                    zipf.write(file_path, arcname=file_path.name)
        response = send_file(zip_path, as_attachment=False)
        response.headers["Content-Disposition"] = build_content_disposition(zip_name, default_stem="batch_download")
        return response

    @app.route("/batch_delete", methods=["POST"])
    def batch_delete():
        payload = request.get_json(silent=True) or {}
        files = payload.get("files", [])
        deleted: List[str] = []
        failed: List[Dict[str, str]] = []
        for item in files:
            file_path = resolve_existing_path(item)
            if not file_path:
                failed.append({"file": str(item), "error": "Invalid path"})
                continue
            try:
                if file_path.exists():
                    file_path.unlink()
                    deleted.append(normalize_file_ref(file_path))
            except Exception as exc:  # noqa: BLE001
                failed.append({"file": str(item), "error": str(exc)})
        scan_files()
        return jsonify({"status": "success", "deleted": deleted, "failed": failed})

    @app.route("/batch_move", methods=["POST"])
    def batch_move():
        payload = request.get_json(silent=True) or {}
        files = payload.get("files", [])
        target_dir = (payload.get("target_dir") or "").strip().replace("\\", "/")
        if not target_dir:
            return jsonify({"status": "error", "message": "Please select target directory"})
        destination_root = (docs_root / target_dir).resolve(strict=False)
        if not is_relative_to(destination_root, docs_root):
            return jsonify({"status": "error", "message": "Invalid target directory"}), 400
        destination_root.mkdir(parents=True, exist_ok=True)

        moved: List[str] = []
        failed: List[Dict[str, str]] = []
        for item in files:
            source = resolve_existing_path(item)
            if not source:
                failed.append({"file": str(item), "error": "Invalid file"})
                continue
            try:
                target = next_available_path(destination_root / source.name)
                shutil.move(str(source), str(target))
                moved.append(normalize_file_ref(target))
            except Exception as exc:  # noqa: BLE001
                failed.append({"file": str(item), "error": str(exc)})
        scan_files()
        return jsonify({"status": "success", "moved": moved, "failed": failed})

    @app.route("/compare_documents", methods=["POST"])
    def compare_documents():
        try:
            doc1 = request.files.get("doc1")
            doc2 = request.files.get("doc2")
            if not doc1 or not doc2:
                return jsonify({"status": "error", "message": "Please upload two files"})
            content1 = upload_to_text(doc1)
            content2 = upload_to_text(doc2)
            lines1 = content1.splitlines()
            lines2 = content2.splitlines()
            diff_html = difflib.HtmlDiff().make_file(lines1, lines2, context=True)
            matcher = difflib.SequenceMatcher(None, lines1, lines2)

            added_lines = 0
            removed_lines = 0
            modified_lines = 0
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == "replace":
                    modified_lines += max(i2 - i1, j2 - j1)
                elif tag == "delete":
                    removed_lines += i2 - i1
                elif tag == "insert":
                    added_lines += j2 - j1
            legal_analysis = analyze_legal_risks(content1, content2, lines1, lines2, matcher)
            return jsonify(
                {
                    "status": "success",
                    "diff_html": diff_html,
                    "stats": {
                        "added_lines": added_lines,
                        "removed_lines": removed_lines,
                        "modified_lines": modified_lines,
                        "total_lines": max(len(lines1), len(lines2)),
                    },
                    "legal_analysis": legal_analysis,
                }
            )
        except Exception as exc:  # noqa: BLE001
            return jsonify({"status": "error", "message": str(exc)})

    @app.route("/recommend_clauses", methods=["POST"])
    def recommend_clauses():
        payload = request.get_json(silent=True) or {}
        case_type = payload.get("case_type", "")
        recommendations = config.clause_templates.get(case_type, config.default_clause_template)
        return jsonify({"status": "success", "recommendations": recommendations})

    @app.route("/create_reminder", methods=["POST"])
    def create_reminder():
        nonlocal reminder_id_counter
        payload = request.get_json(silent=True) or {}
        case_name = payload.get("case_name", "")
        content = payload.get("content", "")
        custom_reminders = payload.get("custom_reminders", [])
        created: List[Dict[str, Any]] = []
        for date_info in extract_dates(content):
            reminder = {
                "id": str(reminder_id_counter),
                "case_name": case_name,
                "event": date_info["event"],
                "reminder_time": date_info["date"],
                "message": f"{case_name} - {date_info['event']} reminder",
                "created_at": datetime.now().isoformat(),
                "status": "pending",
            }
            reminders_db.append(reminder)
            created.append(reminder)
            reminder_id_counter += 1
        for custom in custom_reminders if isinstance(custom_reminders, list) else []:
            reminder_time = normalize_date(str(custom.get("date", "")))
            if not reminder_time:
                continue
            reminder = {
                "id": str(reminder_id_counter),
                "case_name": case_name,
                "event": custom.get("event", "custom reminder"),
                "reminder_time": reminder_time,
                "message": custom.get("message", f"{case_name} - custom reminder"),
                "created_at": datetime.now().isoformat(),
                "status": "pending",
            }
            reminders_db.append(reminder)
            created.append(reminder)
            reminder_id_counter += 1
        return jsonify({"status": "success", "message": f"Created {len(created)} reminders", "reminders": created})

    @app.route("/get_reminders", methods=["GET"])
    def get_reminders():
        today = datetime.now().date()
        due: List[Dict[str, Any]] = []
        for reminder in reminders_db:
            if reminder.get("status") != "pending":
                continue
            date_value = normalize_date(str(reminder.get("reminder_time", "")))
            if not date_value:
                continue
            reminder_date = datetime.strptime(date_value, "%Y-%m-%d").date()
            if reminder_date <= today:
                due.append(reminder)
        return jsonify({"status": "success", "reminders": reminders_db, "due_reminders": due, "total": len(reminders_db)})

    @app.route("/mark_reminder_done", methods=["POST"])
    def mark_reminder_done():
        payload = request.get_json(silent=True) or {}
        reminder_id = payload.get("reminder_id", "")
        for reminder in reminders_db:
            if reminder.get("id") == reminder_id:
                reminder["status"] = "completed"
                break
        return jsonify({"status": "success"})

    @app.route("/delete_reminder", methods=["POST"])
    def delete_reminder():
        payload = request.get_json(silent=True) or {}
        reminder_id = payload.get("reminder_id", "")
        nonlocal reminders_db
        reminders_db = [item for item in reminders_db if item.get("id") != reminder_id]
        return jsonify({"status": "success"})

    return app
